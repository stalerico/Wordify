import os
import random
from flask import Flask, render_template, request, redirect, url_for, send_from_directory, session, flash
from flask_login import LoginManager, login_user, login_required, logout_user, UserMixin, current_user
from flask_migrate import Migrate
from flask_sqlalchemy import SQLAlchemy
from transformers import pipeline
from werkzeug.security import generate_password_hash, check_password_hash
import spacy
from docx import Document

#initialize models
sentiment_analyzer = pipeline("sentiment-analysis", model="cardiffnlp/twitter-roberta-base-sentiment")
token_Model = spacy.load("en_core_web_sm")
token_Model.disable_pipes('ner')

# App Config
app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = './uploads'
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///login.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
app.secret_key = 'welovepink'

# Creating Upload Folder if not exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

# Initializing da Extensions
db = SQLAlchemy(app)
migrate = Migrate(app, db)
login_manager = LoginManager(app)
login_manager.login_view = "login"


with app.app_context():
    db.create_all()

#database models
class User(UserMixin, db.Model):
    __tablename__ = 'login'
    username = db.Column(db.String(80), primary_key=True)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(200), nullable=False)

    def get_id(self):
        return self.username

#utility functions
def user_exists(username, email):
    """Check if user or email already exists in the database."""
    if User.query.filter_by(username=username).first():
        flash("Username already exists, please choose a different one.", 'error')
        return True
    if User.query.filter_by(email=email).first():
        flash("Email is already in use, please choose a different one.", 'error')
        return True
    return False

def currentwordcount(text):
    """Count the words in a given text."""
    return len(text.split())

def updatedwordcount(processed_text):
    """Count words in the processed text."""
    return len(processed_text.split())

def process_file(file):
    """Process uploaded file and extract text."""
    file_extension = os.path.splitext(file.filename)[1]
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], file.filename)
    file.save(file_path)

    if file_extension == ".txt":
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    elif file_extension == ".docx":
        doc = Document(file_path)
        return '\n'.join(paragraph.text for paragraph in doc.paragraphs)
    else:
        return "Unsupported file type. Please upload a .txt or .docx file."

def get_user_by_username(username):
    """Get user from the database by username."""
    return User.query.filter_by(username=username).first()

#routes
#login rmanager
@login_manager.user_loader
def load_user(user_id):
    return User.query.filter_by(username=user_id).first()

#sign in route
@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == "POST":
        username = request.form["username"]
        email = request.form["email"]
        password = request.form["password"]
        hashed_password = generate_password_hash(password)

        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash("Username already exists, please choose a different one.", 'error')
            return redirect(url_for('signup'))

        existing_email = User.query.filter_by(email=email).first()
        if existing_email:
            flash("Email is already in use, please choose a different one.", 'error')
            return redirect(url_for('signup'))

        user = User(username=username, email=email, password=hashed_password)
        db.session.add(user)
        db.session.commit()

        login_user(user)
        flash("Sign-up successful! Welcome to Wordify.", 'success')
        return redirect(url_for('wordify'))

    return render_template('signup.html')

#Login Route
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form["username"]
        password = request.form["password"]

        user = User.query.filter_by(username=username).first()
        if user and check_password_hash(user.password, password):
            login_user(user)
            flash("Login successful! Welcome back.", 'success')
            return redirect(url_for('wordify'))
        else:
            flash("Invalid credentials. Please try again.", 'error')
            return redirect(url_for('login'))

    return render_template('login.html')

#Logout Route
@app.route('/logout')
def logout():
    logout_user()
    flash("You have been logged out.", 'info')
    return redirect(url_for('homepage'))

#account route
@app.route('/account')
@login_required  #only logged-in users can access this
def account():
    username = current_user.username 
    return render_template('account.html', username=username)

#update_info route
@app.route('/update_info', methods=['POST'])
@login_required
def update_info():
    username = request.form['username']
    email = request.form['email']
    password = request.form['password']

    if username != current_user.username:
        existing_user = User.query.filter_by(username=username).first()
        if existing_user:
            flash("Username already exists, please choose a different one.", 'error')
            return redirect(url_for('account'))  #redirect if username is taken

    if email != current_user.email:
        existing_email = User.query.filter_by(email=email).first()
        if existing_email:
            flash("Email is already in use, please choose a different one.", 'error')
            return redirect(url_for('account'))  #redirect back if email is unavaialBLEEE

    if password:
        hashed_password = generate_password_hash(password)
        current_user.password = hashed_password 

    #update username and email (if they are changed)
    current_user.username = username
    current_user.email = email

    db.session.commit()

    flash('Your information has been updated successfully!', 'success')

    return redirect(url_for('account'))

#Wordify Route (protected by login)
@app.route('/wordify', methods=['GET', 'POST'])
@login_required
def wordify():
    text = session.get('text', "")
    processed_text = ""  # Initialize processed text
    current_word_count = session.get('current_word_count', 0)
    updated_word_count = current_word_count

    return render_template('wordify.html', text=text, processed_text=processed_text, 
                           current_word_count=current_word_count, updated_word_count=updated_word_count)

#Upload page route
@app.route('/upload', methods=['POST'])
def upload_file():
    if 'user_file' not in request.files:
        return "No file uploaded."
    
    file = request.files['user_file']
    
    if file.filename == '':
        return "No file selected."

    text = process_file(file)  #Extract text (from the file the user uploaded)
    current_word_count = currentwordcount(text)
    processed_text = ""  # will be updated by the extend button
    
    #store the original text in session(for live editing)
    session['text'] = text
    session['current_word_count'] = current_word_count
    session['filename'] = file.filename
    session['file_extension'] = os.path.splitext(file.filename)[1].lower()

    return render_template('wordify.html', text=text, current_word_count=current_word_count, processed_text=processed_text, filename=session['filename'])

#save text route
@app.route('/save_text', methods=['POST'])
def save_text():
    data = request.get_json()
    updated_text = data.get('text')

    processed_text = process_text(updated_text)  #process updated text
    updated_word_count = updatedwordcount(processed_text)

    return {
        'message': 'Text saved successfully!',
        'processed_text': processed_text,
        'updated_word_count': updated_word_count
    }

#extend text route
@app.route('/extend', methods=['POST'])
def extend_text():
    text = session.get('text', "")
    processed_text = process_text(text)
    updated_word_count = updatedwordcount(processed_text)
    
    session['processed_text'] = processed_text

    file_extension = session.get('file_extension', '.txt')
    processed_filename = f"processed_{session['filename']}"

    #save the processed tetx to allow user to download (for the updated text)
    if file_extension == '.txt':
        processed_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{processed_filename}.txt")
        with open(processed_file_path, 'w', encoding='utf-8') as f:
            f.write(processed_text)
    elif file_extension == '.docx':
        doc = Document()
        doc.add_paragraph(processed_text)
        processed_file_path = os.path.join(app.config['UPLOAD_FOLDER'], f"{processed_filename}.docx")
        doc.save(processed_file_path)

    download_link = url_for('download_file', filename=f"{processed_filename}{file_extension}")
    
    return render_template('wordify.html', text=text, current_word_count=session['current_word_count'], processed_text=processed_text, updated_word_count=updated_word_count, download_link=download_link, filename=session.get('filename'))

#Download file route
@app.route('/uploads/<filename>')
def download_file(filename):
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)

    if not os.path.exists(file_path):
        return "File not found!", 404

    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

#other html routes
@app.route('/')
def homepage():
    return render_template('homepage.html')

@app.route('/contact', methods=['GET', 'POST'])
def contact():
    return render_template('contact.html')

@app.route('/info', methods=['GET', 'POST'])
def info():
    return render_template('info.html')

#MAIN PROCESS TEXT FUNCTION 
def process_text(text):
    """Processes the text using spaCy to expand verbs, adjectives, and nouns, and break down contractions."""
    contractions_w_apos = {"I'm": "I am", "you're": "you are", "he's": "he is", "she's": "she is", "it's": "it is", "we're": "we are", "they're": "they are", "I've": "I have", "you've": "you have", "he's": "he has", "she's": "she has", "it's": "it has", "we've": "we have", "they've": "they have", "I'll": "I will", "you'll": "you will", "he'll": "he will", "she'll": "she will", "it'll": "it will", "we'll": "we will", "they'll": "they will", "I'd": "I would", "you'd": "you would", "he'd": "he would", "she'd": "she would", "it'd": "it would", "we'd": "we would", "they'd": "they would", "don't": "do not", "doesn't": "does not", "didn't": "did not", "isn't": "is not", "aren't": "are not", "wasn't": "was not", "weren't": "were not", "haven't": "have not", "hasn't": "has not", "hadn't": "had not", "won't": "will not", "wouldn't": "would not", "can't": "cannot", "couldn't": "could not", "mustn't": "must not", "should've": "should have", "shan't": "shall not", "shouldn't": "should not", "mightn't": "might not", "needn't": "need not", "oughtn't": "ought not", "let's": "let us", "gonna": "going to", "wanna": "want to", "gotta": "got to", "kinda": "kind of", "sorta": "sort of"}
    contractions_wout_apos = {"Im": "I am", "youre": "you are", "hes": "he is", "shes": "she is", "its": "it is", "theyre": "they are", "Ive": "I have", "youve": "you have", "hes": "he has", "shes": "she has", "its": "it has", "weve": "we have", "theyve": "they have", "Ill": "I will", "youll": "you will", "hell": "he will", "shell": "she will", "itll": "it will", "well": "we will", "theyll": "they will", "Id": "I would", "youd": "you would", "hed": "he would", "shed": "she would", "itd": "it would", "wed": "we would", "theyd": "they would", "dont": "do not", "doesnt": "does not", "didnt": "did not", "isnt": "is not", "arent": "are not", "wasnt": "was not", "werent": "were not", "havent": "have not", "hasnt": "has not", "hadnt": "had not", "wont": "will not", "wouldnt": "would not", "cant": "cannot", "couldnt": "could not", "mustnt": "must not", "shant": "shall not", "shouldve": "should have", "shouldnt": "should not", "mightnt": "might not", "neednt": "need not", "oughtnt": "ought not", "lets": "let us", "gonna": "going to", "wanna": "want to", "gotta": "got to", "kinda": "kind of", "sorta": "sort of"}

    #Expanding contractions
    expanded_text = []
    for word in text.split():
        if word in contractions_w_apos:
            expanded_text.append(contractions_w_apos[word])
        elif word in contractions_wout_apos:
            expanded_text.append(contractions_wout_apos[word])
        else:
            expanded_text.append(word)

    #Joining the words back into a string
    first_iteration_string = ' '.join(expanded_text)
    
    #Processing the text using spaCyyyyy ayyee
    processed_text = token_Model(first_iteration_string)
    altered_text = []

    for token in processed_text:
        word = token.text
        pos = token.pos_
        dependency = token.dep_

        #verbs
        if pos == "VERB" and dependency == "ROOT":
            verb_expander = Verb(word)
            expanded_verb = verb_expander.get_expanded_verb()
            altered_text.append(f'<span class="highlight">{expanded_verb}</span>')

        #adjectives
        elif pos == "ADJ" and dependency in {"amod", "acomp"}:
            adj_expander = Adj(word)
            expanded_adj = adj_expander.get_expanded_adj()
            altered_text.append(f'<span class="highlight">{expanded_adj}</span>')

        #nouns
        elif pos == "NOUN" and dependency in {"nsubj", "dobj", "pobj", "poss"}:
            noun_expander = Noun(word)
            expanded_noun = noun_expander.get_expanded_noun()
            altered_text.append(f'<span class="highlight">{expanded_noun}</span>')

        #everything else
        else:
            altered_text.append(word)
    
    punctuation = {".", ",", "!", "?", ":", ";", "-", "(", ")", "”", "“", "’", "'"}
    reconstructed_text = ''
 
    for word in altered_text:
        if word in punctuation:
            reconstructed_text = reconstructed_text.rstrip() + word
        else:
           
            reconstructed_text += " " + word
 
    reconstructed_text = reconstructed_text.strip()

    return reconstructed_text

#Classes for Adjectives, verbs and nouns and defined methods (to add a preceding descriptor)
class Adj:
    def __init__(self, word):
        self.word = word
        self.descriptors = {
            "positive":["very", "extremely", "incredibly", "absolutely", "so", "highly"],
            "neutral": ["quite", "somewhat", "fairly", "moderately", "slightly"],
            "negative": ["barely", "hardly", "not very", "scarcely", "poorly"]
        }

    def analyze_sentiment(self):
        result = sentiment_analyzer(self.word)
        return result[0]["label"]
    
    def get_expanded_adj(self):
        sentiment = self.analyze_sentiment()

        if sentiment == "LABEL_2":  # Positive
            descriptors = self.descriptors["positive"]
        elif sentiment == "LABEL_1":  # Neutral
            descriptors = self.descriptors["neutral"]
        else:  # Negative
            descriptors = self.descriptors["negative"]

        descriptor = random.choice(descriptors)
        return f"{descriptor} {self.word}"

#Classes for nouns
class Noun:
    def __init__(self, word):
        self.word = word
        self.descriptors= {
            "positive":["amazing", "brilliant", "caring", "charming", "creative", "dedicated", "energetic", "exceptional", "fearless", "generous", "honest", "inspiring", "joyful", "kind", "loyal", "majestic", "optimistic", "passionate", "radiant", "resilient", "reliable", "sincere", "supportive", "talented", "trustworthy", "uplifting", "vibrant", "wonderful", "wise", "zealous"],
            "neutral": ["average", "basic", "common", "normal", "standard", "typical", "usual", "ordinary", "unremarkable", "neutral", "general", "modest", "regular", "mediocre", "slight", "everyday", "typical", "unexceptional", "simple", "routine"],
            "negative": ["angry", "bitter", "broken", "cold", "cruel", "damaged", "deceitful", "depressed", "difficult", "dirty", "disgusting", "dismal", "evil", "fake", "frightening", "harmful", "harsh", "hostile", "ill", "insensitive", "irresponsible", "lazy", "lonely", "nasty", "negative", "nervous", "offensive", "overbearing", "poisonous", "reckless", "rude", "sick", "terrible", "toxic", "useless", "vicious", "weak", "wicked", "worthless"],
            }

    def analyze_sentiment(self):
        result = sentiment_analyzer(self.word)
        return result[0]["label"]
    
    def get_expanded_noun(self):
        sentiment = self.analyze_sentiment()

        if sentiment == "LABEL_2":  # Pos
            descriptors = self.descriptors["positive"]
        elif sentiment == "LABEL_1":  # Neutral
            descriptors = self.descriptors["neutral"]
        else:  # Neg
            descriptors = self.descriptors["negative"]

        descriptor = random.choice(descriptors)
        return f"{descriptor} {self.word}"


#classes for verbs
class Verb:
    def __init__(self, word):
        self.word = word
        self.descriptors= {
            "positive": ["quickly", "eagerly", "enthusiastically", "gracefully", "joyful", "brightly","cheerfully", "vigorously"],
            "negative": ["reluctantly", "slowly", "angrily", "harshly", "sadly", "grudgingly", "heavily","lazily","bitterly"],
            "neutral": ["normally", "quietly", "gently", "steadily", "smoothly", "calmly", "modestly", "softly", "casually"]
            }

    def analyze_sentiment(self):
        result = sentiment_analyzer(self.word)
        return result[0]["label"]

    def get_expanded_verb(self):
        sentiment = self.analyze_sentiment()

        if sentiment == "LABEL_2":  # Pos
            descriptors = self.descriptors["positive"]
        elif sentiment == "LABEL_1":  # Neutral
            descriptors = self.descriptors["neutral"]
        else:  # Neg
            descriptors = self.descriptors["negative"]

        descriptor = random.choice(descriptors)
        return f"{descriptor} {self.word}"

#to debuggg PURRR
if __name__ == '__main__':
    app.run(debug=True)